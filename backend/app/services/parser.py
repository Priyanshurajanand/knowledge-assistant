import os
from typing import List, Dict, Any
from pypdf import PdfReader
from docx import Document as DocxDocument
from app.core.exceptions import DocumentParsingError

class FileParserService:
    @staticmethod
    def parse(file_path: str, filename: str) -> List[Dict[str, Any]]:
        """
        Parses PDF, DOCX, or TXT files.
        Returns a list of dicts: [{'text': str, 'page_number': int}]
        """
        _, ext = os.path.splitext(filename.lower())
        
        try:
            if ext == ".pdf":
                return FileParserService._parse_pdf(file_path)
            elif ext == ".docx":
                return FileParserService._parse_docx(file_path)
            elif ext == ".txt":
                return FileParserService._parse_txt(file_path)
            else:
                raise DocumentParsingError(filename, f"Unsupported file type: {ext}")
        except Exception as e:
            if isinstance(e, DocumentParsingError):
                raise e
            raise DocumentParsingError(filename, str(e))

    @staticmethod
    def _parse_pdf(file_path: str) -> List[Dict[str, Any]]:
        pages = []
        with open(file_path, "rb") as f:
            reader = PdfReader(f)
            for idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                # Clean up null characters and standard whitespace cleanups
                text = text.replace("\x00", "").strip()
                if text:
                    pages.append({
                        "text": text,
                        "page_number": idx + 1
                    })
        return pages

    @staticmethod
    def _parse_docx(file_path: str) -> List[Dict[str, Any]]:
        # python-docx doesn't easily support page divisions, so we treat paragraphs
        # as a stream and group them, or treat the whole document as page 1.
        doc = DocxDocument(file_path)
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
                
        # Also parse tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        text = "\n\n".join(full_text).replace("\x00", "")
        if not text.strip():
            return []
            
        return [{
            "text": text,
            "page_number": 1
        }]

    @staticmethod
    def _parse_txt(file_path: str) -> List[Dict[str, Any]]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read().replace("\x00", "").strip()
        if not text:
            return []
            
        return [{
            "text": text,
            "page_number": 1
        }]
